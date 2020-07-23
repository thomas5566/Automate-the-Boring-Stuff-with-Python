#! python3

import docx

with open('guests.txt') as f:
    names = f.readlines()
    document = docx.Document()

    for name in names:
        name = name.strip()

        document.add_paragraph('It would be a pleasure to have the company of', style='Title')
        document.add_paragraph(name, style='Title')
        document.add_paragraph('at 11010 Memory Lane on the Evening of', style='Title')
        document.add_paragraph('April 1st', style='Title')
        document.add_paragraph("at 7 o'clock", style='Title')

        document.add_page_break()

    document.save('invites.docx')

    print("File has been created and save as'invites.docx'")